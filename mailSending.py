import base64
import logging
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from googleapiclient.discovery import build
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
import os
from docx import Document
import streamlit as st

# Gmail API scope to send email
SCOPES = ['https://www.googleapis.com/auth/gmail.send']
subject = "Reminder to complete your Moodle Courses"

def authenticate_gmail_sender():
    """Authenticate using Streamlit secrets instead of local files."""
    token_info = st.secrets["gcp_service_account"]
    creds = Credentials(
        token=token_info["token"],
        refresh_token=token_info["refresh_token"],
        token_uri=token_info["token_uri"],
        client_id=token_info["client_id"],
        client_secret=token_info["client_secret"],
        scopes=token_info["scopes"]
    )
    return build('gmail', 'v1', credentials=creds)

def generate_body(name: str, incomplete_courses: list[str]) -> str:
    doc = Document(r'downloads\Course Reminder Mail.docx')
    paras = []

    courses = "\n".join(incomplete_courses)

    for para in doc.paragraphs:
        line = para.text.replace("{name}", name).replace("{course_list}", courses)
        paras.append(line)

    text = "\n".join(paras)
    return text

def create_message(sender_email: str, recipient_email: str, name: str, incomplete: list[str]) -> dict:
    """Create a base64-encoded email payload for Gmail API."""
    body = generate_body(name, incomplete)

    message = MIMEMultipart()
    message['To'] = recipient_email
    message['From'] = sender_email
    message['Subject'] = subject
    message.attach(MIMEText(body, 'plain'))

    raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
    return {'raw': raw_message}


def send_email(service, sender_email: str, recipient_email: str, name: str, incomplete: list[str]) -> bool:
    """Send a single email using Gmail API."""
    try:
        message = create_message(sender_email, recipient_email, name, incomplete)
        service.users().messages().send(userId='me', body=message).execute()
        logging.info("Email sent to %s", recipient_email)
        return True
    except Exception as exc:
        logging.error("Failed to send email to %s: %s", recipient_email, exc)
        return False


def send_bulk_emails(sender_email: str, payloads: list[tuple[str, str, list[str]]]) -> list[str]:
    """Send multiple emails using Gmail API."""
    failed: list[str] = []
    service = authenticate_gmail_sender()
    for recipient_email, name, incomplete in payloads:
        success = send_email(service, sender_email, recipient_email, name, incomplete)
        if not success:
            failed.append(recipient_email)
    return failed

def generate_POSH_body(name: str):
    doc = Document(r'downloads\POSH Reminder Mail.docx')
    paras = []
    for para in doc.paragraphs:
        line = para.text.replace("{name}", name)
        paras.append(line)

    body = "\n".join(paras)
    return body

def create_POSH_message(sender_email, recipient_email, recipient_name):
    body = generate_POSH_body(recipient_name)

    message = MIMEMultipart()
    message['To'] = recipient_email
    message['From'] = sender_email
    message['Subject'] = "Reminder to Complete Mandatory POSH Training"
    message.attach(MIMEText(body, 'plain'))

    raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
    return {'raw': raw_message}

def send_POSH_email(service, sender_email: str, recipient_email: str, recipient_name: str) -> bool:
    try:
        message = create_POSH_message(sender_email, recipient_email, recipient_name)
        service.users().messages().send(userId='me', body=message).execute()
        logging.info("Email sent to %s", recipient_email)
        return True
    except Exception as exc:
        logging.error("Failed to send email to %s: %s", recipient_email, exc)
        return False

def send_POSH_emails(sender_email: str, recievers: list[tuple[str, str]]):
    failed: list[str] = []
    service = authenticate_gmail_sender()
    for recipient_name, recipient_email in recievers:
        success = send_POSH_email(service, sender_email, recipient_email, recipient_name)
        if not success:
            failed.append(recipient_email)
    return failed